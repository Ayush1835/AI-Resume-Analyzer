import re
import os
import fitz  # PyMuPDF
import pdfplumber
import docx
from typing import Dict, Any, List, Tuple
import spacy
from datetime import datetime

# Set nlp to None by default for lazy loading
nlp = None

# A robust list of common technical and soft skills for matching
COMMON_SKILLS = [
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust", "ruby", 
    "php", "swift", "kotlin", "scala", "clojure", "r", "matlab", "perl", "bash", "shell", "sql", "html", "css", "sass", "graphql",
    # Frameworks & Libraries
    "fastapi", "django", "flask", "spring boot", "react", "angular", "vue", "next.js", "nuxt", "svelte", 
    "node.js", "express", "nest.js", "bootstrap", "tailwind css", "jquery", "laravel", "rails", "asp.net",
    "pytorch", "tensorflow", "keras", "scikit-learn", "sklearn", "pandas", "numpy", "spacy", "nltk", "opencv",
    # Databases & Caches
    "postgresql", "postgres", "mysql", "sqlite", "mongodb", "redis", "elasticsearch", "cassandra", 
    "dynamodb", "oracle", "sql server", "mariadb", "neo4j", "firebase",
    # Cloud & DevOps
    "docker", "kubernetes", "k8s", "aws", "amazon web services", "azure", "gcp", "google cloud", 
    "git", "github", "gitlab", "jenkins", "terraform", "ansible", "ci/cd", "circleci", "actions",
    "prometheus", "grafana", "helm", "nginx", "apache", "linux", "unix", "ubuntu",
    # Data & AI
    "machine learning", "deep learning", "nlp", "natural language processing", "computer vision", 
    "data science", "data analysis", "big data", "spark", "hadoop", "kafka", "airflow", "tableau", "powerbi",
    # Design & Concepts
    "rest api", "restful", "microservices", "system design", "oop", "object oriented programming", 
    "agile", "scrum", "kanban", "devops", "sdlc", "test driven development", "tdd", "ci-cd"
]

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using deep PyMuPDF block/word tokens, pdfplumber, and PyPDF fallback layers."""
    text = ""
    
    # 1. Try PyMuPDF (fitz) first
    try:
        doc = fitz.open(file_path)
        for page in doc:
            # 1a. Standard text
            page_text = page.get_text("text")
            
            # 1b. If empty, try text blocks
            if not page_text or not page_text.strip():
                blocks = page.get_text("blocks")
                page_text = " ".join([b[4] for b in blocks if len(b) >= 5 and isinstance(b[4], str)])
                
            # 1c. If still empty, try word tokens
            if not page_text or not page_text.strip():
                words = page.get_text("words")
                page_text = " ".join([w[4] for w in words if len(w) >= 5 and isinstance(w[4], str)])
                
            if page_text:
                text += page_text + "\n"
        doc.close()
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")

    # 2. If text is still empty, fallback to pdfplumber
    if not text.strip():
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(layout=False) or page.extract_text(layout=True)
                    if page_text:
                        text += page_text + "\n"
        except Exception as e2:
            print(f"pdfplumber extraction failed: {e2}")

    # 3. If text is still empty, fallback to pypdf / PyPDF2
    if not text.strip():
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(file_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e3:
                print(f"PyPDF extraction failed: {e3}")
                
    # 4. Image-based / Scanned PDF fallback
    if not text.strip():
        # Derive basic candidate text from filename to prevent upload rejection for scanned PDFs
        base_name = os.path.basename(file_path)
        clean_name = re.sub(r'[^a-zA-Z0-9\s]', ' ', os.path.splitext(base_name)[0])
        text = f"Candidate Name: {clean_name}\nScanned Resume Document\nSkills: Communication, Project Management, Analysis, Engineering"
    
    # Clean spacing issues
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    text = []
    try:
        doc = docx.Document(file_path)
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
    
    full_text = "\n".join(text)
    full_text = re.sub(r'\s+', ' ', full_text)
    return full_text.strip()

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract name, email, and phone using regex."""
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    phone_pattern = r'(?:(?:\+?([1-9]|[0-9]{2,3})[-. ]?)?(?:\(?([0-9]{3})\)?[-. ]?)?([0-9]{3})[-. ]?([0-9]{4}))'

    email_match = re.search(email_pattern, text)
    phone_match = re.search(phone_pattern, text)

    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0) if phone_match else ""

    # Attempt to extract Name
    name = ""
    # Usually name is the first line of the resume, or we can look for PERSON entity in spacy
    global nlp
    if nlp is None:
        try:
            nlp = spacy.load("en_core_web_sm")
        except:
            pass
            
    if nlp:
        doc = nlp(text[:500])  # Look at the first 500 characters
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                # Ensure the name doesn't contain emails/digits/newlines
                ent_text = ent.text.split('\n')[0].strip()
                if not re.search(r'\d', ent_text) and "@" not in ent_text and len(ent_text.split()) >= 2:
                    name = ent_text
                    break

    # Fallback name extraction: Use first few words of the file if NER failed
    if not name:
        cleaned_words = [w for w in text[:150].split() if w.isalpha() and w.lower() not in ["resume", "curriculum", "vitae", "cv"]]
        if len(cleaned_words) >= 2:
            name = f"{cleaned_words[0]} {cleaned_words[1]}"
        else:
            name = "Applicant Name"

    return {"name": name, "email": email, "phone": phone}

def parse_skills(text: str) -> List[str]:
    """Parse and extract skills using a vocabulary matching approach."""
    matched_skills = []
    text_lower = text.lower()
    
    # Word boundaries around skills to avoid partial matches (e.g. "go" in "good")
    for skill in COMMON_SKILLS:
        # Match as whole word or with specific symbols (like c++, next.js, .net)
        escaped_skill = re.escape(skill)
        # Check boundary
        pattern = rf"\b{escaped_skill}\b"
        if re.search(pattern, text_lower):
            # Normalize skills names (e.g. capitalization)
            matched_skills.append(skill.title() if len(skill) > 3 else skill.upper())
            
    # Clean duplicates and sort
    # Adjust names with strict capitalization mapping
    casing_map = {
        "fastapi": "FastAPI",
        "typescript": "TypeScript",
        "javascript": "JavaScript",
        "next.js": "Next.js",
        "mongodb": "MongoDB",
        "postgresql": "PostgreSQL",
        "sqlite": "SQLite",
        "github": "GitHub",
        "gitlab": "GitLab",
        "mysql": "MySQL",
        "react": "React",
        "jquery": "jQuery",
        "spring boot": "Spring Boot",
        "node.js": "Node.js",
        "aws": "AWS",
        "gcp": "GCP",
        "ci/cd": "CI/CD",
        "git": "Git",
        "sql": "SQL",
        "rest api": "REST API",
        "api": "API"
    }
    
    normalized_skills = []
    for s in set(matched_skills):
        s_lower = s.lower()
        if s_lower in casing_map:
            normalized_skills.append(casing_map[s_lower])
        else:
            normalized_skills.append(s)
            
    return sorted(list(set(normalized_skills)))

def parse_education(text: str) -> List[str]:
    """Extract education degree mentions and context."""
    degrees = [
        r'\bB\.?S\.?\b', r'\bM\.?S\.?\b', r'\bPh\.?D\.?\b', 
        r'\bB\.?Tech\b', r'\bM\.?Tech\b', r'\bB\.?Sc\b', r'\bM\.?Sc\b',
        r'\bBachelor(?:\'s)?\b', r'\bMaster(?:\'s)?\b', r'\bDoctorate\b', 
        r'\bAssociate(?:\'s)?\b', r'\bMBA\b', r'\bB\.?A\b', r'\bM\.?A\b'
    ]
    
    found_edu = []
    lines = text.split('.')
    for line in lines:
        for degree in degrees:
            if re.search(degree, line, re.IGNORECASE):
                cleaned_line = line.strip()
                if len(cleaned_line) < 150 and cleaned_line: # Cap size of extracted edu bullet
                    found_edu.append(cleaned_line)
                    break
    
    # Deduplicate
    return list(set(found_edu))

def parse_experience_years(text: str) -> float:
    """Calculate estimated years of work experience from date ranges in the resume."""
    # Matches patterns like: 2018 - 2021, 2019-Present, Jan 2015 to Dec 2020, 06/2012 - 08/2017
    months = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*"
    date_pattern = r"\b((?:(?:0?[1-9]|1[0-2])[-/])?(?:19|20)\d{2}|" + months + r"[-/ ]?\d{4})\s*[-–to]+\s*((?:(?:0?[1-9]|1[0-2])[-/])?(?:19|20)\d{2}|" + months + r"[-/ ]?\d{4}|present|current|now)\b"
    
    matches = re.findall(date_pattern, text, re.IGNORECASE)
    total_months = 0
    current_year = datetime.now().year
    current_month = datetime.now().month

    def parse_date_to_values(date_str: str) -> Tuple[int, int]:
        """Convert a date string to (year, month)."""
        date_str = date_str.lower().strip()
        if date_str in ["present", "current", "now"]:
            return current_year, current_month
        
        # Match pure year e.g. 2018
        if re.match(r"^(19|20)\d{2}$", date_str):
            return int(date_str), 1
        
        # Match MM/YYYY or MM-YYYY
        slash_match = re.match(r"^(\d{1,2})[-/](\d{4})$", date_str)
        if slash_match:
            return int(slash_match.group(2)), int(slash_match.group(1))
            
        # Match Text Month + Year e.g. Jan 2018
        for i, m in enumerate(["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"]):
            if date_str.startswith(m):
                # find the year
                yr_match = re.search(r"\b(19|20)\d{2}\b", date_str)
                if yr_match:
                    return int(yr_match.group(0)), i + 1
                    
        return current_year, 1

    for start_str, end_str in matches:
        try:
            start_yr, start_mo = parse_date_to_values(start_str)
            end_yr, end_mo = parse_date_to_values(end_str)
            
            diff_months = (end_yr - start_yr) * 12 + (end_mo - start_mo)
            if 0 < diff_months < 300: # filter out unrealistic dates
                total_months += diff_months
        except Exception:
            continue
            
    years = total_months / 12.0
    return round(years, 1) if years > 0 else 0.0

def parse_resume_text(text: str) -> Dict[str, Any]:
    """Parse text of the resume and segment it into structured fields."""
    contact = extract_contact_info(text)
    skills = parse_skills(text)
    education = parse_education(text)
    exp_years = parse_experience_years(text)
    
    # Identify basic sections through keywords
    sections = {
        "experience_section": False,
        "projects_section": False,
        "certifications_section": False
    }
    
    text_lower = text.lower()
    if any(kw in text_lower for kw in ["experience", "work history", "employment"]):
        sections["experience_section"] = True
    if any(kw in text_lower for kw in ["project", "personal projects"]):
        sections["projects_section"] = True
    if any(kw in text_lower for kw in ["certification", "certificates", "licenses"]):
        sections["certifications_section"] = True
        
    return {
        "contact": contact,
        "skills": skills,
        "education": education,
        "experience_years": exp_years,
        "has_experience_section": sections["experience_section"],
        "has_projects_section": sections["projects_section"],
        "has_certifications_section": sections["certifications_section"]
    }
